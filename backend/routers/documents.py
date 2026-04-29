import io
import copy
from datetime import datetime
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document

from backend import models, auth as auth_service
from backend.dependencies import get_db

router = APIRouter(prefix="/api/technician", tags=["documents"])

TEMPLATE_PATH = Path("docs_oficiales/CSS_RG-10. Oferta Ed.05_ES.docx")


def _replace_in_runs(paragraph, placeholder: str, value: str):
    """Replace a placeholder across potentially split runs in a paragraph."""
    full_text = paragraph.text
    if placeholder not in full_text:
        return

    # Find which runs contain the placeholder (may span multiple runs)
    runs = paragraph.runs
    combined = ""
    for run in runs:
        combined += run.text

    if placeholder not in combined:
        return

    # Strategy: rebuild runs preserving formatting of the first run that touches the placeholder
    new_text = combined.replace(placeholder, value)

    # Clear all runs and set the full text in the first run, preserving its formatting
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""


def _build_services_text(services):
    """Build a readable text block listing all active services."""
    lines = []
    total = 0.0
    for s in services:
        if s.is_deleted:
            continue
        price = s.quoted_price or 0.0
        total += price
        line = f"• {s.service_name} — {s.hours}h"
        if s.quoted_price is not None:
            line += f" — {s.quoted_price:.2f} €"
        if s.comment:
            line += f"  ({s.comment})"
        lines.append(line)
    lines.append(f"\nTOTAL: {total:.2f} €")
    return "\n".join(lines)


@router.get("/offers/{offer_id}/document")
def generate_offer_document(
    offer_id: int,
    current_user=Depends(auth_service.require_technician_or_higher),
    db: Session = Depends(get_db),
):
    """Generate a pre-filled Word document (.docx) from the official MiNa template."""

    if not TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="Document template not found on server")

    offer = db.query(models.Offer).filter(models.Offer.id == offer_id).first()
    if not offer:
        raise HTTPException(status_code=404, detail="Offer not found")

    client = offer.client
    manager = offer.manager

    # Build replacement map
    replacements = {
        "QUOTE": offer.reference or str(offer.id),
        "TECHNICIAN": f"{manager.first_name} {manager.last_name}" if manager else "Not assigned",
        "DATE": offer.created_at.strftime("%d/%m/%Y") if offer.created_at else datetime.now().strftime("%d/%m/%Y"),
        "CLIENT": f"{client.first_name} {client.last_name}" if client else "Unknown",
        "PROJECT_CODE": client.codigo_proyecto or "—" if client else "—",
        "CCII": client.cuenta_interna or "—" if client else "—",
        "IIPP": client.investigador_principal or "—" if client else "—",
        "TITLE": f"Offer {offer.reference or offer.id}",
        "TABLE": _build_services_text(offer.services),
        "DELIVERY": "To be confirmed",
    }

    # Load and fill the template
    doc = Document(str(TEMPLATE_PATH))

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            _replace_in_runs(paragraph, placeholder, value)

    # Also check inside tables (if the template has any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        _replace_in_runs(paragraph, placeholder, value)

    # Stream the filled document back
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Oferta_{offer.reference or offer.id}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
